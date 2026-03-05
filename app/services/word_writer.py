from __future__ import annotations

import re
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Literal, Tuple

from docx import Document

from app.services.excel_reader import MaintenanceRow


# Accept both normal '}}' and typo ']}' endings found in some templates.
PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([^{}\[\]]+?)\s*(?:\}\}|\]\})")

TemplateMode = Literal["row_expansion", "single_cell_multirow"]


@dataclass
class CrossValidationResult:
    excel_count: int
    word_count: int
    mismatches: List[str]

    @property
    def is_match(self) -> bool:
        return self.excel_count == self.word_count and not self.mismatches


@dataclass
class TemplateLayout:
    mode: TemplateMode
    table_index: int
    placeholder_row_index: int | None = None
    target_cell: Tuple[int, int] | None = None


def _clone_row_and_cell_props(src_row, dst_row) -> None:
    """Clone row/cell formatting (especially borders) from src_row to dst_row."""
    try:
        src_trPr = src_row._tr.get_or_add_trPr()
        dst_trPr = dst_row._tr.get_or_add_trPr()
        for child in list(dst_trPr):
            dst_trPr.remove(child)
        for child in list(src_trPr):
            dst_trPr.append(deepcopy(child))
    except Exception:
        pass

    for i in range(min(len(src_row.cells), len(dst_row.cells))):
        src_cell = src_row.cells[i]
        dst_cell = dst_row.cells[i]
        try:
            src_tcPr = src_cell._tc.get_or_add_tcPr()
            dst_tcPr = dst_cell._tc.get_or_add_tcPr()
            for child in list(dst_tcPr):
                dst_tcPr.remove(child)
            for child in list(src_tcPr):
                dst_tcPr.append(deepcopy(child))
        except Exception:
            pass


def _extract_placeholders(text: str) -> List[str]:
    if not text:
        return []
    return [m.group(1).strip() for m in PLACEHOLDER_PATTERN.finditer(text) if m.group(1).strip()]


def _render_template_text(template_text: str, data: Dict[str, str]) -> str:
    if not template_text:
        return ""

    def replacer(match: re.Match[str]) -> str:
        field = match.group(1).strip()
        return data.get(field, "")

    return PLACEHOLDER_PATTERN.sub(replacer, template_text)


def _set_cell_text(cell, text: str) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.add_run(text)


def _iter_placeholder_cells(doc: Document, table_index: int) -> List[Tuple[int, int, str, List[str]]]:
    table = doc.tables[table_index]
    items: List[Tuple[int, int, str, List[str]]] = []
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            fields = _extract_placeholders(cell.text)
            if fields:
                items.append((row_idx, col_idx, cell.text, fields))
    return items


def _detect_template_layout(doc: Document) -> TemplateLayout:
    if not doc.tables:
        raise ValueError("No tables found in Word template.")

    table_index = len(doc.tables) - 1
    cells = _iter_placeholder_cells(doc, table_index)
    if not cells:
        raise ValueError("No {{field}} placeholders found in template.")

    # Single-cell multi-row layout (VPN template): one placeholder cell with repeated fields.
    if len(cells) == 1:
        row_idx, col_idx, _, fields = cells[0]
        if len(fields) > len(set(fields)):
            return TemplateLayout(
                mode="single_cell_multirow",
                table_index=table_index,
                target_cell=(row_idx, col_idx),
            )

    # Default layout: placeholder row is expanded as many rows as Excel records.
    placeholder_row_index = min(row_idx for row_idx, _, _, _ in cells)
    return TemplateLayout(
        mode="row_expansion",
        table_index=table_index,
        placeholder_row_index=placeholder_row_index,
    )


def _extract_repeating_unit_template(text: str) -> str:
    matches = list(PLACEHOLDER_PATTERN.finditer(text))
    if not matches:
        return text.strip()

    first_token = text[matches[0].start() : matches[0].end()]
    second_pos = text.find(first_token, matches[0].end())
    if second_pos == -1:
        return text.strip()

    # Remove separators accidentally included before the next repeated token.
    return text[:second_pos].rstrip(" \t\r\n[")


def extract_template_fields(template_path: Path) -> List[str]:
    """Extract unique placeholder field names from the whole template."""
    if template_path.suffix.lower() != ".docx":
        raise ValueError("Word template must be .docx")

    doc = Document(str(template_path))
    if not doc.tables:
        raise ValueError("No tables found in Word template.")

    fields: List[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for field in _extract_placeholders(cell.text):
                    if field not in fields:
                        fields.append(field)

    if not fields:
        raise ValueError("No {{field}} placeholders found in template.")

    return fields


def _fill_row_expansion(table, placeholder_row_index: int, rows: List[MaintenanceRow]) -> None:
    placeholder_row = table.rows[placeholder_row_index]
    cell_templates: List[str] = [cell.text for cell in placeholder_row.cells]
    if not any(_extract_placeholders(text) for text in cell_templates):
        raise ValueError("No {{field}} placeholders found in template data row.")

    while len(table.rows) > placeholder_row_index + 1:
        table._tbl.remove(table.rows[placeholder_row_index + 1]._tr)

    def fill_row(table_row, data: Dict[str, str]) -> None:
        for idx, template_text in enumerate(cell_templates):
            if idx >= len(table_row.cells):
                break
            _set_cell_text(table_row.cells[idx], _render_template_text(template_text, data))

    fill_row(placeholder_row, rows[0])

    for item in rows[1:]:
        new_row = table.add_row()
        _clone_row_and_cell_props(placeholder_row, new_row)
        fill_row(new_row, item)


def _fill_single_cell_multirow(table, target_cell: Tuple[int, int], rows: List[MaintenanceRow]) -> None:
    target_row, target_col = target_cell
    cell = table.cell(target_row, target_col)
    unit_template = _extract_repeating_unit_template(cell.text)
    rendered = [_render_template_text(unit_template, row).strip() for row in rows]
    combined = "\n\n".join(block for block in rendered if block)
    _set_cell_text(cell, combined)

    # Replace other one-time placeholders (if any) using the first row.
    first_row = rows[0]
    for row_idx, row in enumerate(table.rows):
        for col_idx, cur in enumerate(row.cells):
            if (row_idx, col_idx) == target_cell:
                continue
            if _extract_placeholders(cur.text):
                _set_cell_text(cur, _render_template_text(cur.text, first_row))


def fill_template_and_save(template_path: Path, rows: List[MaintenanceRow]) -> Path:
    """Fill Word template table and save output."""
    if template_path.suffix.lower() != ".docx":
        raise ValueError("Word template must be .docx")
    if not rows:
        raise ValueError("No valid data rows found in Excel file.")

    doc = Document(str(template_path))
    layout = _detect_template_layout(doc)
    table = doc.tables[layout.table_index]

    if layout.mode == "row_expansion":
        if layout.placeholder_row_index is None:
            raise ValueError("Template placeholder row was not detected.")
        _fill_row_expansion(table, layout.placeholder_row_index, rows)
    else:
        if layout.target_cell is None:
            raise ValueError("Template target cell was not detected.")
        _fill_single_cell_multirow(table, layout.target_cell, rows)

    out_base = template_path.with_suffix("").name + "_작성완료"
    out_docx = template_path.parent / f"{out_base}.docx"
    doc.save(str(out_docx))
    return out_docx


def _validate_row_expansion(
    template_doc: Document, output_doc: Document, layout: TemplateLayout, rows: List[MaintenanceRow]
) -> CrossValidationResult:
    row_idx = layout.placeholder_row_index
    if row_idx is None:
        raise ValueError("Template placeholder row was not detected.")

    template_table = template_doc.tables[layout.table_index]
    output_table = output_doc.tables[layout.table_index]
    if row_idx >= len(template_table.rows):
        raise ValueError("Template placeholder row index is out of range.")
    if row_idx >= len(output_table.rows):
        raise ValueError("Generated placeholder row index is out of range.")

    cell_templates: List[str] = [cell.text for cell in template_table.rows[row_idx].cells]
    output_data_rows = output_table.rows[row_idx:]

    mismatches: List[str] = []
    max_rows = min(len(rows), len(output_data_rows))
    max_cols = len(cell_templates)

    for r in range(max_rows):
        expected_row = rows[r]
        word_row = output_data_rows[r]
        if len(word_row.cells) < max_cols:
            mismatches.append(
                f"Row {r + 1}: Word has fewer cells ({len(word_row.cells)}) than expected ({max_cols})."
            )
            continue
        for c in range(max_cols):
            expected = _render_template_text(cell_templates[c], expected_row).strip()
            actual = word_row.cells[c].text.strip()
            if expected != actual:
                mismatches.append(f"Row {r + 1}, Col {c + 1}: expected '{expected}', got '{actual}'.")

    return CrossValidationResult(excel_count=len(rows), word_count=len(output_data_rows), mismatches=mismatches)


def _validate_single_cell_multirow(
    template_doc: Document, output_doc: Document, layout: TemplateLayout, rows: List[MaintenanceRow]
) -> CrossValidationResult:
    target_cell = layout.target_cell
    if target_cell is None:
        raise ValueError("Template target cell was not detected.")

    table_idx = layout.table_index
    t_row, t_col = target_cell
    template_table = template_doc.tables[table_idx]
    output_table = output_doc.tables[table_idx]

    template_text = template_table.cell(t_row, t_col).text
    unit_template = _extract_repeating_unit_template(template_text)
    expected_text = "\n\n".join(
        _render_template_text(unit_template, row).strip() for row in rows if _render_template_text(unit_template, row).strip()
    ).strip()
    actual_text = output_table.cell(t_row, t_col).text.strip()

    mismatches: List[str] = []
    if expected_text != actual_text:
        mismatches.append("Single-cell multi-row content mismatch in target work-history cell.")

    # Validate additional one-time placeholders with the first row values.
    first_row = rows[0]
    for row_idx, row in enumerate(template_table.rows):
        for col_idx, cell in enumerate(row.cells):
            if (row_idx, col_idx) == target_cell:
                continue
            if not _extract_placeholders(cell.text):
                continue
            expected = _render_template_text(cell.text, first_row).strip()
            actual = output_table.cell(row_idx, col_idx).text.strip()
            if expected != actual:
                mismatches.append(
                    f"One-time cell mismatch at ({row_idx + 1}, {col_idx + 1}): expected '{expected}', got '{actual}'."
                )

    return CrossValidationResult(excel_count=len(rows), word_count=len(rows), mismatches=mismatches)


def validate_generated_report(
    template_path: Path, output_path: Path, rows: List[MaintenanceRow]
) -> CrossValidationResult:
    """Cross-check generated Word data against Excel-derived rows."""
    template_doc = Document(str(template_path))
    output_doc = Document(str(output_path))
    if not template_doc.tables:
        raise ValueError("No tables found in Word template.")
    if not output_doc.tables:
        raise ValueError("No tables found in generated Word document.")
    if not rows:
        raise ValueError("No valid data rows found in Excel file.")

    layout = _detect_template_layout(template_doc)
    if layout.mode == "row_expansion":
        return _validate_row_expansion(template_doc, output_doc, layout, rows)
    return _validate_single_cell_multirow(template_doc, output_doc, layout, rows)
